#Import important libraries
from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
import os

pdf_lines = []
worddoc_lines = []

#define appender for adding pdf files
def append_pdf(input,output):
    [output.addPage(input.getPage(page_num)) for page_num in range(input.numPages)]

# define appender for adding word documents
def combine_word_documents(files):
    combined_document = Document()
    count, number_of_files = 0, len(files)
    for file in files:
        source_document = Document(file)

        # Don't add a page break if you've
        # reached the last file.
        for paragraph in source_document.paragraphs:
            text = paragraph.text
            combined_document.add_paragraph(text)

    combined_document.save('combined_word_documents.docx')


#check if input file is present, for names of pdf
if(os.path.isfile("pdf_list.txt")):
    text_file = open("pdf_list.txt", "r")
    lines = text_file.read().split(',')
    text_file.close()
else:
    print('pdf_list.txt is missing')

#check if names are present in input file 
if(len(lines)==0):
    print('No pdf file names - empty.')

    
if(os.path.isfile("result.pdf")):
    os.remove("result.pdf")
if(os.path.isfile("combined_word_documents.docx")):
    os.remove("combined_word_documents.docx")



#concatenate pdf files together into single input 

for line in lines:
    if ".pdf" in line:
        pdf_lines.append(line)
    elif ".docx" in line:
        worddoc_lines.append(line)
    else:
        print("this file is not supported : " + line)

print("Atleast it started")
output = PdfFileWriter()
for line in pdf_lines:
    append_pdf(PdfFileReader((open(str(line),"rb")),strict=False),output)
output.write(open("result.pdf","wb"))
print("Pdf concatenated successfully into: result.pdf")


#combine_word_documents(worddoc_lines)
print("Word doc concatenated successfully into: combined_word_documents.docx")


